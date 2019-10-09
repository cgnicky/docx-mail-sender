import re
from docx import Document
from datetime import timedelta
from datetime import date
import os
from Crypto.Cipher import AES
import base64
from pyhocon import ConfigFactory
from send_mail import send_gmail


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def check_and_save(file):
    print("Checking existing file...")
    if os.path.isfile(file) is True:
        print("Removing file...")
        os.remove(file)
        print("Saving the file to {}".format(file))
        doc.save(file)
    else:
        print("Saving the file to {}".format(file))
        doc.save(file)


def decode_passcode(encoded_msg, key):
    padding = '{'
    cipher = AES.new(key.encode())
    decode_aes = lambda c, e: str(c.decrypt(base64.b64decode(e)), 'utf-8').rstrip(padding)
    result = decode_aes(cipher, encoded_msg)

    return result


# Loading type-safe configuration
conf = ConfigFactory.parse_file("application.conf")

# Current datetime
dt = date.today()
# Datetime of next Monday
start_dt = dt + timedelta(days=-dt.weekday(), weeks=1)
# Datetime of next Friday
end_dt = start_dt + timedelta(days=4)

#Reason of minus 1911 is because of Taiwan's year
apply_year = str(int(dt.strftime("%Y")) - 1911)
apply_month = dt.strftime("%m")
apply_day = dt.strftime("%d")

start_year = str(int(start_dt.strftime("%Y")) - 1911)
start_month = start_dt.strftime("%m")
start_day = start_dt.strftime("%d")

end_year = str(int(end_dt.strftime("%Y")) - 1911)
end_month = end_dt.strftime("%m")
end_day = end_dt.strftime("%d")

sub_start_date = start_dt.strftime("%Y_%m%d")
sub_end_date = end_dt.strftime("%m%d")

subject_text = (conf.get_string("mail_setup.subject_title")).format(sub_start_date, sub_end_date)

mail_content = (conf.get_string("mail_setup.content_text")).format(start_dt.strftime("%Y%m%d"),
                                                                   end_dt.strftime("%Y%m%d"))

dictionary = {'st_yyy': start_year, 'st_m': start_month, 'st_dd': start_day, 'et_yyy': end_year,
              'et_m': end_month, 'et_dd': end_day, 'apply_yyy': apply_year, 'apply_m': apply_month,
              'apply_dd': apply_day, 'work_desc': conf.get_string("doc_setup.work_desc")}

sender_mail = conf.get_string("mail_setup.sender_addr")
receiver_mail = conf.get_string("mail_setup.receiver_addr")
encrypted_passcode = conf.get_string("mail_setup.password").encode()
passkey = conf.get_string("mail_setup.passkey")
pdf_to_attach = conf.get_list("mail_setup.attachment_path")

doc_to_be_save = (conf.get_string("doc_setup.save_path")).format(start_dt.strftime("%Y%m%d"), end_dt.strftime("%Y%m%d"))
original_doc = conf.get_string("doc_setup.template_path")
doc = Document(original_doc)
for word, replacement in dictionary.items():
    word_re = re.compile(word)
    docx_replace_regex(doc, word_re, replacement)

# docx_replace_regex(doc, regex1 , replace1)
check_and_save(doc_to_be_save)
send_gmail(doc_to_be_save, pdf_to_attach, sender_mail, receiver_mail, decode_passcode(encrypted_passcode, passkey),
           subject_text, mail_content)
